import datetime
from io import BytesIO
import os
import re
import time
import uuid
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls  # Fix namespace lỗi
import requests
from app.lib.string import (
    split_toktak_url,
)
from app.lib.logger import logger
from app.makers.images import ImageMaker
import zipfile

from app.lib.header import generate_desktop_user_agent

date_create = datetime.datetime.now().strftime("%Y_%m_%d")
UPLOAD_FOLDER = os.path.join(os.getcwd(), f"uploads/{date_create}")
CURRENT_DOMAIN = os.environ.get("CURRENT_DOMAIN") or "http://localhost:5000"


class DocxMaker:
    def make(self, title, ads_text, description, images=[], batch_id=0):
        timestamp = int(time.time())
        unique_id = uuid.uuid4().hex
        file_name = f"{timestamp}_{unique_id}.docx"

        docx_path = f"{UPLOAD_FOLDER}/{batch_id}/{file_name}"

        doc = Document()

        doc.add_heading(title, level=1)

        if ads_text != "":
            notice_para = doc.add_paragraph()
            notice_run = notice_para.add_run(ads_text)
            notice_run.font.size = Pt(8)

        user_agent = generate_desktop_user_agent()
        headers = {
            "User-Agent": user_agent,
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7",
            "Accept-Encoding": "gzip, deflate, br, zstd",
            "Accept-Language": "en,vi;q=0.9,es;q=0.8,vi-VN;q=0.7,fr-FR;q=0.6,fr;q=0.5,en-US;q=0.4",
        }

        for item in description:
            if item.startswith("IMAGE_URL_"):
                index = int(item.split("_")[-1])
                if 0 <= index < len(images):
                    image_url = images[index]
                    try:
                        response = requests.get(image_url, headers=headers)
                        if response.status_code == 200:
                            image_stream = BytesIO(response.content)
                            doc.add_picture(
                                image_stream, width=Inches(6.5)
                            )  # Full width
                    except Exception as e:
                        print(f"⚠️ Lỗi tải ảnh: {e}")
            else:
                doc.add_paragraph(item)

        doc.save(docx_path)

        file_size = os.path.getsize(docx_path)
        mime_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        docx_url = f"{CURRENT_DOMAIN}/files/{date_create}/{batch_id}/{file_name}"
        return {
            "file_size": file_size,
            "mime_type": mime_type,
            "docx_url": docx_url,
        }

    def make_txt(self, title, ads_text, description, images=[], batch_id=0):
        timestamp = int(time.time())
        unique_id = uuid.uuid4().hex
        file_name = f"{timestamp}_{unique_id}"
        txt_path = f"{UPLOAD_FOLDER}/{batch_id}/{file_name}.txt"
        zip_path = f"{UPLOAD_FOLDER}/{batch_id}/{file_name}.zip"
        os.makedirs(os.path.dirname(txt_path), exist_ok=True)
        txt_lines = []

        txt_lines.append(title + "\n\n")

        if ads_text:
            txt_lines.append(ads_text + "\n\n")

        for item in description:
            if item.startswith("IMAGE_URL_"):
                txt_lines.append("\n\n\n\n")
            else:
                split_lines = split_toktak_url(item)
                for line in split_lines:
                    line = re.sub(r"<br\s*/?>", "", line)
                    txt_lines.append(line + "\n")
                txt_lines.append("\n")

        with open(txt_path, "w", encoding="utf-8") as f:
            f.writelines(txt_lines)

        image_downloads = []
        check_image = {}
        for image_url in images:
            image_path = ImageMaker.save_image_url_get_path(image_url, batch_id)
            if not image_path:
                continue
            if "toktak.ai" in image_url:
                check_image[image_path] = True
            else:
                check_image[image_path] = False
            image_downloads.append(image_path)

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            # 🔍 Thêm file TXT nếu tồn tại
            if os.path.isfile(txt_path):
                zipf.write(txt_path, arcname=os.path.basename(txt_path))

            for img_path in image_downloads:
                if os.path.isfile(img_path):
                    zipf.write(img_path, arcname=os.path.basename(img_path))

        try:
            if os.path.isfile(txt_path):
                os.remove(txt_path)

            for img_path in image_downloads:
                if check_image.get(img_path):
                    continue
                if os.path.isfile(img_path):
                    os.remove(img_path)

        except Exception as e:
            logger.error(f"⚠️ Lỗi khi xoá file sau khi zip: {e}")

        docx_url = f"{CURRENT_DOMAIN}/files/{date_create}/{batch_id}/{file_name}.zip"
        file_size = os.path.getsize(zip_path)
        return {
            "zip_path": zip_path,
            "file_size": file_size,
            "docx_url": docx_url,
            "mime_type": "application/zip",
        }

    def add_hyperlink(self, paragraph, url, text=None):
        """Chèn hyperlink vào đoạn văn bản"""
        text = text or url  # Nếu không có text, dùng URL làm text

        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        # XML chuẩn (Loại bỏ nsdecls gây lỗi)
        hyperlink_xml = f"""
        <w:hyperlink r:id="{r_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
            <w:r>
                <w:rPr>
                    <w:color w:val="0000FF"/>
                    <w:u w:val="single"/>
                </w:rPr>
                <w:t>{text}</w:t>
            </w:r>
        </w:hyperlink>
        """

        # Thêm Hyperlink vào đoạn văn bản
        hyperlink = parse_xml(hyperlink_xml)
        paragraph._element.append(hyperlink)
